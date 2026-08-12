"""Reading back a .docx this project itself rendered.

Fixtures here are produced by ``render_document`` rather than hand-built or
checked in as bytes. That is the whole point of the pairing: the preview exists
to show the reader what a Task actually produced, so the thing it is tested
against has to be what a Task actually produces. A hand-made document would let
the renderer change its styles -- which it does, for Chinese thesis format --
and leave this suite green while the console showed a wall of unstyled text.

Two kinds of fixture cannot come from there, and both say so where they are
built. ``_zip_of`` assembles archives the renderer would never produce, because
refusing a bomb is the one thing it cannot be asked to demonstrate. The counts
of what a preview drops need the opposite: a document with a picture, a
footnote or a floating figure, none of which ``render_document`` can make -- it
renders headings, prose, bullets and tables. Those are built with python-docx
directly, and so is the one case that needs a document *poorer* than a rendered
one: every document this project renders carries a header, a footer and a
custom style, so a test asking whether anything at all was reported missing
would pass on those three whatever happened to the table it dropped. The
control for each of them is still a rendered document, so the
assertion that matters most -- that an ordinary Task output reports no pictures
and no footnotes -- is made against what a Task actually produces.
"""

from __future__ import annotations

import io
import os
import re
import struct
import zipfile
import zlib
from dataclasses import fields

import pytest
from docx import Document
from docx.document import Document as WordDocument
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.text.run import Run

from agent_workbench.adapters.documents.docx import (
    MAX_PREVIEW_CHARS,
    MAX_PREVIEW_COMPRESSION_RATIO,
    MAX_PREVIEW_EXPANDED_BYTES,
    DocxPreview,
    DocxTooLargeError,
    extract_docx_preview,
    preflight_docx,
)
from agent_workbench.apps.word_mcp.contract import (
    DocumentRequest,
    DocumentSection,
    SimpleTable,
)
from agent_workbench.apps.word_mcp.renderer import render_document


def _document(*sections: DocumentSection, title: str = "季度报告") -> bytes:
    return render_document(
        DocumentRequest(title=title, subtitle=None, sections=sections)
    )


def _section(
    heading: str,
    *paragraphs: str,
    bullets: tuple[str, ...] = (),
    table: SimpleTable | None = None,
) -> DocumentSection:
    return DocumentSection(
        heading=heading,
        paragraphs=paragraphs,
        bullets=bullets,
        table=table,
    )


def test_a_rendered_document_reads_back_as_its_own_text() -> None:
    preview = extract_docx_preview(
        _document(_section("背景", "第一段正文。", "第二段正文。"))
    )

    assert "背景" in preview.text
    assert "第一段正文。" in preview.text
    assert "第二段正文。" in preview.text
    assert preview.truncated is False


def test_paragraphs_keep_the_order_the_document_has_them_in() -> None:
    preview = extract_docx_preview(
        _document(
            _section("第一节", "甲"),
            _section("第二节", "乙"),
        )
    )

    assert preview.text.index("甲") < preview.text.index("乙")


def test_a_table_is_read_in_place_rather_than_after_the_prose() -> None:
    """The bug this walks the body's XML children to avoid.

    ``document.paragraphs`` and ``document.tables`` each return their own kind
    in order, so reading one then the other puts every table after every
    paragraph -- a report whose table sits in section 1 would show it below the
    conclusion. The control is the paragraph that follows it.
    """

    preview = extract_docx_preview(
        _document(
            _section(
                "数据",
                "表格之前。",
                table=SimpleTable(headers=("指标", "值"), rows=(("营收", "12"),)),
            ),
            _section("结论", "表格之后。"),
        )
    )

    # Located by the rendered *table row*, not by its cell text. Cell text also
    # exists as a paragraph inside the table, so `index("指标")` finds that copy
    # -- which sits in the right place even when the traversal is wrong, and
    # made an earlier version of this assertion survive the exact bug it names.
    lines = preview.text.splitlines()
    table_at = next(i for i, line in enumerate(lines) if line.startswith("| 指标 |"))
    before_at = next(i for i, line in enumerate(lines) if line == "表格之前。")
    after_at = next(i for i, line in enumerate(lines) if line == "表格之后。")

    assert before_at < table_at < after_at
    assert preview.table_count == 1


def test_a_table_becomes_a_markdown_table_the_console_can_render() -> None:
    preview = extract_docx_preview(
        _document(
            _section(
                "数据",
                table=SimpleTable(
                    headers=("指标", "值"),
                    rows=(("营收", "12"), ("成本", "8")),
                ),
            )
        )
    )

    assert "| 指标 | 值 |" in preview.text
    assert "| --- | --- |" in preview.text
    assert "| 营收 | 12 |" in preview.text


def test_a_pipe_inside_a_cell_does_not_break_the_row_it_sits_in() -> None:
    preview = extract_docx_preview(
        _document(
            _section(
                "数据",
                table=SimpleTable(headers=("键", "值"), rows=(("a|b", "1"),)),
            )
        )
    )

    assert r"a\|b" in preview.text
    # Escaped, so the row still has exactly the two cells it was given. Split
    # on unescaped pipes -- which is what a Markdown renderer does, and the
    # thing an unescaped cell would break.
    row = next(line for line in preview.text.splitlines() if "a\\|b" in line)
    cells = [cell for cell in re.split(r"(?<!\\)\|", row) if cell.strip()]
    assert cells == [" a\\|b ", " 1 "]


def test_a_long_document_is_cut_and_says_so() -> None:
    """Truncation is reported, never silent.

    A preview that stops early and presents itself as complete is worse than no
    preview: the reader draws conclusions from a document they have only part
    of, with nothing on screen to suggest it.
    """

    long_paragraph = "长" * 2_000
    preview = extract_docx_preview(
        _document(_section("很长的一节", *([long_paragraph] * 40)))
    )

    assert preview.truncated is True
    assert len(preview.text) <= MAX_PREVIEW_CHARS + len(long_paragraph)


def test_a_document_that_fits_is_not_marked_truncated() -> None:
    """The control for the test above. Without it, an extractor that reported
    `truncated=True` unconditionally would pass."""

    preview = extract_docx_preview(_document(_section("短", "一句话。")))

    assert preview.truncated is False


def test_a_document_with_no_tables_counts_none() -> None:
    """The control for `table_count`, which the console uses to tell the reader
    what the preview left out."""

    preview = extract_docx_preview(_document(_section("背景", "只有正文。")))

    assert preview.table_count == 0


def _saved(document: WordDocument) -> bytes:
    """A python-docx document as package bytes, for the fixtures below.

    Everything reached through this helper is a document the renderer cannot
    produce -- pictures, footnote marks, numbering it cancels, a package with
    no running titles at all. The rule this file states at the top holds for
    everything that has an alternative; these do not.
    """

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _png() -> bytes:
    """A 1x1 greyscale PNG, built here rather than checked in as a file.

    python-docx reads the header to size the picture, so the bytes have to be a
    real PNG rather than a placeholder. Nothing downstream looks at the pixel:
    the count reads the drawing element the picture produces, not the image.
    """

    def chunk(kind: bytes, payload: bytes) -> bytes:
        body = kind + payload
        return (
            struct.pack(">I", len(payload)) + body + struct.pack(">I", zlib.crc32(body))
        )

    # Width, height, bit depth, colour type 0 (greyscale), and the three
    # compression/filter/interlace bytes a plain PNG fixes at zero.
    header = struct.pack(">IIBBBBB", 1, 1, 8, 0, 0, 0, 0)
    # One scanline: the row's filter byte, then its single sample.
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", header)
        + chunk(b"IDAT", zlib.compress(b"\x00\x00"))
        + chunk(b"IEND", b"")
    )


def test_a_picture_is_counted_because_the_prose_around_it_reads_as_complete() -> None:
    """The omission a reader has no way to infer.

    Truncated text announces itself by stopping mid-sentence. A dropped figure
    leaves paragraphs that read as a finished argument, so the count is the
    only thing on screen that can raise the question at all -- which is why the
    surrounding text being intact is asserted here rather than assumed.
    """

    document = Document()
    document.add_paragraph("图之前的一段。")
    document.add_paragraph("图注。").add_run().add_picture(io.BytesIO(_png()))
    document.add_paragraph("图之后的一段。")

    preview = extract_docx_preview(_saved(document))

    assert preview.image_count == 1
    assert "图之前的一段。" in preview.text
    assert "图之后的一段。" in preview.text


def test_a_rendered_document_reports_no_pictures() -> None:
    """The control, and the one that decides whether the count is worth having:
    a Task's own output has no pictures in it, so a counter that found some
    would park a permanent phantom in the panel beside every run."""

    preview = extract_docx_preview(_document(_section("背景", "只有正文。")))

    assert preview.image_count == 0


def _floated(run: Run) -> None:
    """Move this run's inline picture under an anchor, in place.

    Built by moving a real picture rather than by writing anchor XML, because
    what these cases have to tell apart is a picture from the element around
    it -- a hand-written anchor would be the fixture asserting itself.
    python-docx places only inline pictures, and this is the edit Word makes
    when a reader turns one into a floating figure: the same `a:graphic`,
    carrying the same `pic:pic` and its relationship to the image part, under
    `wp:anchor` instead of `wp:inline`.

    The anchor is left without the position and wrap children Word writes
    beside the graphic. Those decide where on the page the picture lands, which
    is layout, and nothing here reads them.
    """

    drawing = run._r.find(qn("w:drawing"))
    inline = drawing.find(qn("wp:inline"))
    anchor = parse_xml(f"<wp:anchor {nsdecls('wp')}/>")
    for child in list(inline):
        anchor.append(child)
    drawing.replace(inline, anchor)


def test_a_floating_picture_is_counted_although_inline_shapes_cannot_see_it() -> None:
    """Why this count is not `document.inline_shapes`.

    That property's XPath matches `wp:inline` under a run and nothing else, so
    a document whose figures are anchored -- how anybody places an image they
    want text to wrap around -- reports none at all, and on screen "no
    pictures" is indistinguishable from "none of the kind I looked for". Both
    readings are asserted, so the difference is the test rather than a claim
    about it.

    The fixture used to be an empty `wp:anchor`, which is the same element and
    not the same claim: it passes against a count that looks for pictures and
    against one that looks for DrawingML containers, so it could not tell those
    two apart -- and the containers hold text boxes and charts as readily as
    figures. It moved to the case below, as a thing that is not a picture.
    """

    document = Document()
    run = document.add_paragraph("环绕排版的图。").add_run()
    run.add_picture(io.BytesIO(_png()))
    _floated(run)
    content = _saved(document)

    # A real picture, floating: asserted rather than assumed, because the case
    # is about what is inside the anchor and not about the anchor.
    body = Document(io.BytesIO(content)).element.body
    assert len(body.xpath(".//wp:anchor")) == 1
    assert len(body.xpath(".//pic:pic")) == 1
    assert len(Document(io.BytesIO(content)).inline_shapes) == 0
    assert extract_docx_preview(content).image_count == 1


#: Where Word keeps a text box. A Microsoft extension rather than an ECMA-376
#: namespace, so it is absent from python-docx's `nsmap` and declared here.
_WPS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"

#: A text box: a shape whose body happens to hold paragraphs. This is what Word
#: writes for one, and it arrives in the same container a figure does.
_TEXT_BOX = (
    f'<a:graphicData uri="{_WPS}"><wps:wsp xmlns:wps="{_WPS}">'
    '<wps:cNvSpPr txBox="1"/><wps:spPr/><wps:txbx><w:txbxContent>'
    "<w:p><w:r><w:t>框里的字</w:t></w:r></w:p>"
    "</w:txbxContent></wps:txbx><wps:bodyPr/></wps:wsp></a:graphicData>"
)

#: A chart. The chart itself lives in its own part, so what the body holds is
#: the reference to it -- here without the relationship id, which nothing in
#: this module reads and no fixture in this file could satisfy.
_CHART = (
    '<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/chart">'
    "<c:chart/></a:graphicData>"
)

#: SmartArt, the same way.
_SMART_ART = (
    '<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/diagram">'
    "<dgm:relIds/></a:graphicData>"
)


def _drawing(graphic: str, *, floating: bool) -> str:
    """One `w:drawing` holding this graphic, anchored or inline.

    Both containers are covered because a picture arrives in either one, so a
    count narrowed to only the anchored one would still report an inline text
    box -- which Word writes whenever the box sits in the run of text rather
    than beside it -- as a picture.
    """

    container = "wp:anchor" if floating else "wp:inline"
    return (
        f"<w:drawing {nsdecls('w', 'wp', 'a', 'c', 'dgm')}>"
        f'<{container}><wp:extent cx="914400" cy="914400"/>'
        f'<wp:docPr id="1" name="图形 1"/>'
        f"<a:graphic>{graphic}</a:graphic></{container}></w:drawing>"
    )


@pytest.mark.parametrize(
    "drawing",
    [
        pytest.param(_drawing(_TEXT_BOX, floating=True), id="floating-text-box"),
        pytest.param(_drawing(_TEXT_BOX, floating=False), id="inline-text-box"),
        pytest.param(_drawing(_CHART, floating=True), id="chart"),
        pytest.param(_drawing(_SMART_ART, floating=True), id="smart-art"),
        pytest.param(
            f"<w:drawing {nsdecls('w', 'wp')}><wp:anchor/></w:drawing>",
            id="empty-anchor",
        ),
    ],
)
def test_a_drawing_that_holds_no_picture_is_not_counted_as_one(drawing: str) -> None:
    """`wp:inline` and `wp:anchor` are containers, not pictures.

    DrawingML puts text boxes, charts, SmartArt and every drawn shape in the
    same two elements a figure arrives in, so a count of the containers reports
    furniture as photographs. The module used to count exactly that, while the
    comment beside it excluded VML because `w:pict` "also carries text boxes
    and drawn lines" -- word for word what these two do, in a module that was
    counting them anyway. Both halves moved: the count reads `pic:pic`, and the
    reason given for leaving VML out is now one that survives being applied to
    both formats.

    What the reader is told is 图片没有显示 · 1 张, of a document that holds no
    picture. That row sits in the panel's list of what the preview lost, so a
    phantom in it costs the rows beside it -- the footnotes and the headers
    that really are missing -- the credibility the list exists to have.

    The empty anchor is the degenerate member rather than a document Word
    writes. It is here because it is what the floating-picture test above used
    to offer as a picture.
    """

    document = Document()
    document.add_paragraph("图形之前的一段。")
    run = document.add_paragraph("图形之后的一段。").add_run()
    run._r.append(parse_xml(drawing))
    content = _saved(document)

    # A drawing, and no picture in it. Asserted so that a case which stopped
    # being either fails rather than passes for a reason nobody chose.
    body = Document(io.BytesIO(content)).element.body
    assert len(body.xpath(".//w:drawing")) == 1
    assert body.xpath(".//pic:pic") == []

    assert extract_docx_preview(content).image_count == 0


def test_only_the_picture_is_counted_when_a_text_box_sits_beside_it() -> None:
    """The pair, so neither half can be bought by giving up the other.

    A count that went back to `document.inline_shapes` to stop counting text
    boxes says 0 here, and has lost the floating figure `wp:anchor` was added
    for; one that keeps counting containers says 2. Only a count that asks what
    the container holds says 1, which is what the document has.
    """

    document = Document()
    picture = document.add_paragraph("环绕排版的图。").add_run()
    picture.add_picture(io.BytesIO(_png()))
    _floated(picture)
    shape = document.add_paragraph("旁边的文本框。").add_run()
    shape._r.append(parse_xml(_drawing(_TEXT_BOX, floating=True)))

    assert extract_docx_preview(_saved(document)).image_count == 1


def test_a_rendered_document_says_its_letterhead_did_not_come_through() -> None:
    """Headers and footers live in their own parts, which this reader does not
    open -- so the count is the only trace of them, asserted here alongside the
    absence it stands for."""

    preview = extract_docx_preview(_document(_section("背景", "只有正文。")))

    assert preview.header_count == 1
    assert preview.footer_count == 1
    assert "AGENT WORKBENCH" not in preview.text


def test_a_document_with_no_running_titles_counts_none() -> None:
    """The control for the pair above. Every document this project renders
    carries one header and one footer, so without a fixture that has neither, a
    counter that returned a constant 1 would pass."""

    document = Document()
    document.add_paragraph("没有页眉页脚。")

    preview = extract_docx_preview(_saved(document))

    assert preview.header_count == 0
    assert preview.footer_count == 0


def test_a_list_is_counted_because_its_markers_are_not_in_the_text() -> None:
    """Word generates list markers at layout time out of `numbering.xml`, so
    the text of an item carries no sign that it was one. An ordered procedure
    previews as loose lines in the right order, which reads as prose."""

    preview = extract_docx_preview(
        _document(_section("步骤", bullets=("先做这个。", "再做那个。")))
    )

    assert preview.numbered_paragraph_count == 2
    # Nothing of the marker reached the line, which is what leaves the count as
    # the only evidence the list was one.
    item = next(line for line in preview.text.splitlines() if "先做这个。" in line)
    assert item == "先做这个。"


def test_a_document_without_a_list_counts_no_numbering() -> None:
    """The control for the count above."""

    preview = extract_docx_preview(_document(_section("背景", "一段散文。")))

    assert preview.numbered_paragraph_count == 0


def test_a_paragraph_that_cancels_its_numbering_is_not_counted_as_a_list() -> None:
    """`w:numId` of 0 is not a list.

    OOXML spends that value on the opposite meaning: it is how a paragraph
    turns off the numbering its style would otherwise give it. Both paragraphs
    are in the same fixture so that the assertion tells the filter apart from a
    counter that simply found nothing.
    """

    document = Document()
    for num_id in ("1", "0"):
        paragraph = document.add_paragraph(f"编号 {num_id}")
        paragraph._p.get_or_add_pPr().append(
            parse_xml(
                f"<w:numPr {nsdecls('w')}>"
                f'<w:ilvl w:val="0"/><w:numId w:val="{num_id}"/>'
                f"</w:numPr>"
            )
        )

    preview = extract_docx_preview(_saved(document))

    assert preview.numbered_paragraph_count == 1


def test_a_footnote_mark_is_counted_although_its_text_is_in_another_part() -> None:
    """What the reader loses is not the superscript but the sentence it pointed
    at, which is usually the qualification on the claim beside it.

    The mark is injected: python-docx has no footnote API at all, in either
    direction. That costs this fixture less than it looks like it should --
    the note's own part is one this reader never opens, so the mark in the body
    is both what a real document has here and the whole of what is read.
    """

    document = Document()
    run = document.add_paragraph("一句有保留的断言。").add_run()
    run._r.append(parse_xml(f'<w:footnoteReference {nsdecls("w")} w:id="2"/>'))

    preview = extract_docx_preview(_saved(document))

    assert preview.footnote_count == 1
    assert "一句有保留的断言。" in preview.text


def test_a_rendered_document_reports_no_footnotes() -> None:
    """The control for the count above."""

    preview = extract_docx_preview(_document(_section("背景", "只有正文。")))

    assert preview.footnote_count == 0


def test_the_styles_this_project_renders_are_counted_as_flattened() -> None:
    """The limit ADR-043 §7 asked to be moved out of a code comment.

    Headings are recognised by built-in style name, so `AW Title` and
    `AW Bullet` -- what this project's own renderer applies -- arrive as bare
    lines. The document's own title previewing as plain text is asserted
    directly beside the count: it is the visible half of the same fact, and a
    count that kept moving while the output quietly gained a "# " would be
    measuring something else.
    """

    preview = extract_docx_preview(
        _document(_section("背景", "正文。", bullets=("要点一。", "要点二。")))
    )

    # One `AW Title` -- the document's title -- and two `AW Bullet`.
    assert preview.flattened_paragraph_count == 3
    assert "季度报告" in preview.text
    assert "# 季度报告" not in preview.text


def test_a_heading_this_reader_can_express_is_not_counted_as_flattened() -> None:
    """The control that gives the count its meaning.

    Without it, an implementation that counted every paragraph not styled
    "Normal" would pass the test above -- and would report the headings it
    rendered perfectly well as structure it had lost.
    """

    document = Document()
    document.add_heading("标题", level=0)  # style "Title"
    document.add_heading("第一章", level=1)  # style "Heading 1"

    preview = extract_docx_preview(_saved(document))

    assert preview.flattened_paragraph_count == 0
    assert "# 标题" in preview.text
    assert "# 第一章" in preview.text


def test_ordinary_body_text_is_not_counted_as_flattened() -> None:
    """The other control. A number that grew with the document's length would
    be a word count wearing a different name, and a reader told that twenty
    paragraphs were flattened would go looking for structure that was never
    there."""

    document = Document()
    for index in range(20):
        document.add_paragraph(f"第 {index} 段。")

    preview = extract_docx_preview(_saved(document))

    assert preview.flattened_paragraph_count == 0


def test_the_counts_are_of_the_document_and_not_of_the_part_that_fits() -> None:
    """The failure that would make every count above misleading.

    Counting during the walk that builds the text ties each number to wherever
    that walk stopped, and the reader is then shown "truncated" beside "one
    picture" -- from which the only available conclusion is that the rest is
    prose. The picture, the list item and the footnote here all sit past the
    cut, and none of their text reaches the preview.
    """

    document = Document()
    for _ in range(40):
        document.add_paragraph("长" * 2_000)
    document.add_paragraph("尾部图注。").add_run().add_picture(io.BytesIO(_png()))
    numbered = document.add_paragraph("尾部列表项。")
    numbered._p.get_or_add_pPr().append(
        parse_xml(
            f'<w:numPr {nsdecls("w")}><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr>'
        )
    )
    cited = document.add_paragraph("尾部脚注句。").add_run()
    cited._r.append(parse_xml(f'<w:footnoteReference {nsdecls("w")} w:id="2"/>'))

    preview = extract_docx_preview(_saved(document))

    assert preview.truncated is True
    # None of the tail is in the text ...
    assert "尾部图注。" not in preview.text
    assert "尾部列表项。" not in preview.text
    assert "尾部脚注句。" not in preview.text
    # ... and all of it is in the counts.
    assert preview.image_count == 1
    assert preview.numbered_paragraph_count == 1
    assert preview.footnote_count == 1


def test_the_table_count_is_of_the_document_like_every_count_beside_it() -> None:
    """The reversal of `test_the_table_count_alone_stops_where_the_text_does`.

    That test pinned the opposite of this -- `cut.table_count == 0` for a
    document holding one -- and pinned it deliberately: the number was already
    on screen, ADR-045 §6 recorded the seam as a known inconsistency, and the
    pin existed so that moving it later could not happen by accident. This is
    that move, and what overturns the earlier judgement is the test below:
    the seam is not a narrower claim about tables, it is the one path by which
    the panel reports *nothing missing* from a preview that dropped a table
    whole. "Do not change a number the console already shows" is an argument
    about continuity, and it has nothing to say to a number that is wrong.

    What the console gives up is narrower than it sounds. The two readings
    agree on every preview that was not cut, so the only display this changes
    is the truncated one -- where the old number was the wrong answer.

    The same bytes are read twice, at the preview's ceiling and with none, so
    the difference is the ceiling and nothing else.
    """

    content = _document(
        _section("很长的一节", *(["长" * 2_000] * 40)),
        _section(
            "数据", table=SimpleTable(headers=("指标", "值"), rows=(("营收", "12"),))
        ),
    )

    cut = extract_docx_preview(content)
    whole = extract_docx_preview(content, max_chars=None)

    assert cut.truncated is True
    assert whole.truncated is False
    # The table is below the cut, so the text has it one way and not the other
    # -- which is what makes the counts below a statement about the counts.
    assert "| 指标 |" in whole.text
    assert "| 指标 |" not in cut.text
    # And it is in the count either way, like every count beside it.
    assert cut.table_count == whole.table_count == 1
    assert cut.header_count == whole.header_count == 1
    assert cut.flattened_paragraph_count == whole.flattened_paragraph_count == 1


def _reported_counts(preview: DocxPreview) -> dict[str, int]:
    """Every number this preview reports, read off the result rather than listed.

    Read, because which field carries "a table was dropped" is not decided
    here -- `table_count` widened to the whole document, or a second number
    beside it -- and a list written today would name the wrong one either way.
    What the panel does with them is the same for all of them: `PreviewGaps`
    takes every count, drops the zeros, and renders whatever is left.

    `truncated` is left out although `bool` is an `int`, and the exclusion is
    the point rather than a technicality. That flag is already on screen, and
    it is what turns an empty list into a false statement rather than a
    redundant one: a reader told 这里只显示开头 and shown nothing missing
    concludes that what stopped was prose.
    """

    values: dict[str, object] = {
        field.name: getattr(preview, field.name) for field in fields(preview)
    }
    return {
        name: value
        for name, value in values.items()
        if isinstance(value, int) and not isinstance(value, bool)
    }


def test_a_dropped_table_is_never_reported_as_a_faithful_preview() -> None:
    """The one state the panel must not be able to reach: something is missing
    and nothing is listed.

    `PreviewGaps` renders one row per non-zero count and returns null when all
    of them are zero, and its own comment says that the empty list *is* the
    statement that the preview is faithful. `table_count` is the single number
    of what the walk reached rather than of the document, so a table below the
    cut subtracts itself from the only count that would have mentioned it. The
    document loses a table; the panel says nothing was lost.

    Truncation being on screen does not cover this. The console's 这里只显示开头
    says the text stops, and a reader also shown an empty gap list can only
    conclude that what stopped was prose -- which is the inference the module's
    own docstring gives as the reason every other count is of the whole
    document.

    Deliberately not asserted against a field name: whether the fix widens
    `table_count` or adds a number beside it is the next decision, and what is
    pinned here is only that the reader can see *something*. Today's answer is
    pinned directly above, in
    `test_the_table_count_alone_stops_where_the_text_does`; one of the two has
    to move, and this is the one that says which way.
    """

    document = Document()
    for _ in range(40):
        document.add_paragraph("正" * 2_000)
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "被丢掉的表格"
    content = _saved(document)

    whole = extract_docx_preview(content, max_chars=None)
    cut = extract_docx_preview(content)

    # This document has exactly one thing it can lose, which is what makes the
    # assertion below about the table rather than about whatever else was in
    # the fixture. A rendered document would arrive with a header, a footer and
    # a custom style, and pass on those three alone.
    assert whole.table_count == 1
    assert whole.image_count == whole.footnote_count == 0
    assert whole.header_count == whole.footer_count == 0
    assert whole.numbered_paragraph_count == whole.flattened_paragraph_count == 0

    # At the preview's ceiling it loses it -- the table and every word in it.
    assert cut.truncated is True
    assert "被丢掉的表格" not in cut.text

    counts = _reported_counts(cut)
    assert any(count > 0 for count in counts.values()), (
        f"a table is gone from the text and from every number beside it: {counts}"
    )


def test_a_preview_that_lost_nothing_reports_nothing_missing() -> None:
    """The control, and the reason the assertion above asks for *some* number
    rather than a particular one.

    An empty gap list has to stay reachable. A count that is non-zero whenever
    a preview exists would satisfy the case above and put a row under every
    faithful preview in the console -- the same false statement pointed the
    other way, and the rows that mean something would be read as the furniture
    it had just made them look like.
    """

    document = Document()
    document.add_paragraph("只有一句话。")

    preview = extract_docx_preview(_saved(document))

    assert preview.truncated is False
    assert set(_reported_counts(preview).values()) == {0}


def _zip_of(members: dict[str, bytes]) -> bytes:
    """A deflated archive built from raw members, for the refusals below.

    Hand-built rather than rendered, and that is the one place in this file
    where it has to be: the point of these cases is a package the renderer
    would never produce, so a fixture that went through it could not express
    them.
    """

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, payload in members.items():
            archive.writestr(name, payload)
    return buffer.getvalue()


def test_a_package_that_expands_past_the_ceiling_is_refused_before_parsing() -> None:
    """The gap the compressed-size ceiling never closed.

    A .docx is a zip, so what reading one costs is its *expanded* size, while
    the caller's limit is on the stored bytes. This archive is well inside any
    stored-size limit and asks for far more than that once opened; the refusal
    has to happen before python-docx is handed the bytes, because by then the
    allocation has already been made.

    Built to sit *under* the ratio limit on purpose -- a megabyte of
    incompressible data pads the stored size -- so that what it exercises is
    the absolute ceiling and not the ratio check below it.
    """

    padding = os.urandom(1024 * 1024)
    bomb = _zip_of(
        {
            "word/media/noise.bin": padding,
            "word/document.xml": b"\0" * (150 * 1024 * 1024),
        }
    )
    expanded = len(padding) + 150 * 1024 * 1024

    assert expanded > MAX_PREVIEW_EXPANDED_BYTES
    assert expanded < len(bomb) * MAX_PREVIEW_COMPRESSION_RATIO
    with pytest.raises(DocxTooLargeError):
        preflight_docx(bomb)


def test_the_preview_itself_refuses_a_bomb_and_not_only_the_preflight() -> None:
    """The wiring, which is the half that actually protects the route.

    A preflight nothing calls is a function, not a ceiling. Every caller of
    `extract_docx_preview` has to inherit it -- including the ones that do not
    know it exists.
    """

    bomb = _zip_of({"word/document.xml": b"\0" * (400 * 1024 * 1024)})

    with pytest.raises(DocxTooLargeError):
        extract_docx_preview(bomb)


def test_a_small_package_that_is_nothing_but_expansion_is_refused() -> None:
    """Under the absolute ceiling, over the ratio.

    Split across several members on purpose: a ratio measured per entry would
    let one bomb be delivered as ten, so it is measured against the whole
    stored object.
    """

    piece = b"\0" * (6 * 1024 * 1024)
    spread = _zip_of({f"word/part{index}.xml": piece for index in range(6)})
    expanded = 6 * len(piece)

    assert expanded < MAX_PREVIEW_EXPANDED_BYTES
    assert expanded > len(spread) * MAX_PREVIEW_COMPRESSION_RATIO
    with pytest.raises(DocxTooLargeError):
        preflight_docx(spread)


def test_a_package_with_too_many_entries_is_refused() -> None:
    many = _zip_of({f"word/media/image{index}.bin": b"x" for index in range(600)})

    with pytest.raises(DocxTooLargeError):
        preflight_docx(many)


def test_understating_the_declared_size_yields_no_extra_bytes() -> None:
    """Why reading the declared sizes is enough, asserted rather than assumed.

    The preflight trusts ``ZipInfo.file_size``, which the archive supplies
    about itself, so the whole design rests on a bomb being unable to
    understate it. This rewrites the declared size to zero and pins what
    ``zipfile`` then does: it stops the member at the declared length and fails
    the CRC. python-docx reads through the same ``zipfile``, so it is bounded
    by the same number the preflight checked -- no extra bytes are reachable by
    lying, only an unreadable file.
    """

    bomb = bytearray(_zip_of({"word/document.xml": b"\0" * (64 * 1024 * 1024)}))
    # The uncompressed size sits in both the local header and the central
    # directory record; zero it wherever the real value appears.
    real = (64 * 1024 * 1024).to_bytes(4, "little")
    assert real in bomb
    patched = bytes(bomb).replace(real, (0).to_bytes(4, "little"))

    with (
        zipfile.ZipFile(io.BytesIO(patched)) as archive,
        pytest.raises(zipfile.BadZipFile),
    ):
        archive.read("word/document.xml")

    # And the preview says "unreadable", not "here is your document". The
    # exact type is python-docx's business; what matters is that nothing is
    # returned as a readable document.
    with pytest.raises((ValueError, KeyError, zipfile.BadZipFile)):
        extract_docx_preview(patched)


def test_an_ordinary_rendered_document_passes_the_preflight() -> None:
    """The control. Without it, a preflight that refused everything would pass
    every test above, and the preview would be gone rather than bounded."""

    document = _document(
        _section("背景", "第一段正文。" * 200),
        _section("结论", "结论段落。" * 200),
    )

    preflight_docx(document)
    assert extract_docx_preview(document).text != ""


# --------------------------------------------------------------------------
# The styles a document did not define
# --------------------------------------------------------------------------


def _word_styled(*styled: tuple[str, str]) -> bytes:
    """A document wearing Word's own style names, which no fixture above does.

    Everything else in this file is either ``render_document``'s output or a
    python-docx default, and both wear "Normal". That is exactly the blind spot
    ADR-043 §8 asked for samples against: a document somebody wrote in Word
    wears built-in styles this project never applies.
    """

    document = Document()
    for text, style in styled:
        document.add_paragraph(text, style=style)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.mark.parametrize(
    "style",
    ["List Paragraph", "Quote", "Caption", "Body Text 2", "Intense Quote"],
)
def test_a_built_in_style_is_not_reported_as_one_this_document_defined(
    style: str,
) -> None:
    """Word's own styles are not this document's, so they are not counted.

    The count names a specific loss -- a style the document defined and the
    preview could not express. A built-in is dropped too, and that is the
    acknowledged cost of narrowing (see ``_count_flattened_paragraphs``); what
    is not acceptable is reporting Word's furniture as the author's intent.
    """

    preview = extract_docx_preview(_word_styled(("一段正文。", style)))

    assert preview.flattened_paragraph_count == 0


def test_a_list_item_is_reported_once_and_not_twice() -> None:
    """One paragraph, one number.

    "List Paragraph" is what Word puts on every list item. Counting it as a
    flattened style too produced a panel that said twenty paragraphs lost their
    numbering *and* twenty lost their style, of the same twenty paragraphs.
    """

    document = Document()
    for index in range(3):
        paragraph = document.add_paragraph(f"第 {index} 项", style="List Paragraph")
        paragraph._p.get_or_add_pPr().append(
            parse_xml(
                f"<w:numPr {nsdecls('w')}>"
                '<w:ilvl w:val="0"/><w:numId w:val="1"/>'
                "</w:numPr>"
            )
        )
    buffer = io.BytesIO()
    document.save(buffer)

    preview = extract_docx_preview(buffer.getvalue())

    assert preview.numbered_paragraph_count == 3
    assert preview.flattened_paragraph_count == 0


def test_a_style_the_document_defined_is_still_counted() -> None:
    """The control for the narrowing above.

    Without it, a count that had simply stopped working would pass every test
    in this section. This project's own renderer applies custom styles, and
    those are precisely what the field exists to report.
    """

    preview = extract_docx_preview(_document(_section("背景", "第一段正文。")))

    assert preview.flattened_paragraph_count > 0
