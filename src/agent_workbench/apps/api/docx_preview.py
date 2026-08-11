"""Reading a .docx back as text, so the UI can show one instead of describing it.

A rendered document was the one Task output the console could not display. Text
and JSON previewed inline; a .docx -- the artifact most Tasks are actually asked
for -- rendered as "这个类型只能下载查看", which tells the reader everything
about the file except what it says.

Extraction lives on the server rather than in the browser, and the reason is not
convenience. A .docx is a zip of XML, so previewing one client-side means
shipping a zip reader and an XML parser to every page load in order to re-derive
text this process can already produce -- ``python-docx`` is a core dependency
because the Word MCP server renders with it. The same library reads.

**This is a preview, not a conversion.** It recovers the document's text in
reading order and drops everything that made it a document: styles, images,
headers, page geometry. That is the correct trade for a panel beside the run --
somebody who needs the document downloads it, and the download is unchanged
bytes. What this must never do is look like more than it is, which is why the
response says how much it dropped rather than presenting a silent excerpt as the
whole file.
"""

from __future__ import annotations

import io
import zipfile
from collections.abc import Iterator
from dataclasses import dataclass
from typing import Any, Final, cast

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

#: Bounded for the same reason every other preview here is: this is one panel
#: beside a run, and a reader who wants a 200-page report wants the file.
MAX_PREVIEW_CHARS: Final[int] = 40_000

#: Entries in the package. A .docx has on the order of ten to thirty; a couple
#: of hundred means an image-heavy one. Far past that is not a document, and
#: refusing on the count costs nothing because the count is in the central
#: directory -- no member has to be decompressed to read it.
MAX_PREVIEW_ZIP_ENTRIES: Final[int] = 512

#: What the package may weigh once opened. The caller's ceiling is on the
#: *compressed* bytes, which is not the cost of reading one: a .docx is a zip,
#: python-docx expands the whole package, and XML compresses roughly ten to
#: twenty times. So a file inside a 20 MiB limit can still ask this process for
#: a multi-hundred-megabyte allocation, which is the gap this closes.
MAX_PREVIEW_EXPANDED_BYTES: Final[int] = 100 * 1024 * 1024

#: Expanded-to-stored ratio, refused above this. Catches the small file that is
#: nothing but expansion -- a few hundred kilobytes that unpack into tens of
#: megabytes pass the absolute ceiling above while still being a bomb. Set well
#: clear of ordinary documents: prose XML lands around ten to twenty, and the
#: thesis documents this project renders measure in that range.
MAX_PREVIEW_COMPRESSION_RATIO: Final[int] = 200


class DocxTooLargeError(ValueError):
    """The package would cost more to open than a preview may spend.

    Separate from a parse failure because the two mean different things to a
    caller: a document that will not parse is a fact about the file, while this
    is a refusal to try. The route maps them to different statuses.
    """


def preflight_docx(content: bytes) -> None:
    """Refuse a package whose expansion this process should not attempt.

    Runs before ``Document`` sees the bytes, because by then it is too late --
    python-docx expands the package to parse it, so a check that ran after it
    would be reporting an allocation that already happened.

    **Why the declared sizes are enough here.** ``ZipInfo.file_size`` is read
    out of the archive's own central directory, so on its face it is a claim
    the file makes about itself, and the obvious worry is a bomb that simply
    understates it. It cannot: ``ZipExtFile`` stops a member at exactly
    ``file_size`` bytes and then fails the CRC, so an understated declaration
    yields ``BadZipFile`` rather than more bytes. python-docx reads through the
    same ``zipfile``, so whatever bounds this bounds it. An earlier draft of
    this function decompressed every member to measure it instead; that pass
    could not observe anything the declaration did not already permit, and
    charged every legitimate preview a full extra expansion for it.
    """

    try:
        archive = zipfile.ZipFile(io.BytesIO(content))
    except zipfile.BadZipFile as error:
        # Not a size refusal: this one is a parse failure, and is raised as the
        # kind of error the caller already handles for unreadable content.
        raise ValueError("the file is not a readable .docx package") from error

    with archive:
        entries = archive.infolist()
        if len(entries) > MAX_PREVIEW_ZIP_ENTRIES:
            raise DocxTooLargeError(
                f"the package holds {len(entries)} entries, "
                f"more than the {MAX_PREVIEW_ZIP_ENTRIES} a preview will open"
            )
        expanded = sum(entry.file_size for entry in entries)

    if expanded > MAX_PREVIEW_EXPANDED_BYTES:
        raise DocxTooLargeError(
            f"the package expands to {expanded} bytes, "
            f"more than the {MAX_PREVIEW_EXPANDED_BYTES} a preview will hold"
        )
    # Against the whole stored object rather than per member, so that splitting
    # one bomb across several entries does not evade it.
    if content and expanded > len(content) * MAX_PREVIEW_COMPRESSION_RATIO:
        raise DocxTooLargeError(
            f"the package expands {expanded // len(content)}x, "
            f"past the {MAX_PREVIEW_COMPRESSION_RATIO}x a preview will open"
        )


#: How a heading becomes Markdown. python-docx reports built-in heading styles
#: as "Heading 1".."Heading 9"; anything else is a paragraph, including the
#: custom styles the project's own renderer applies for Chinese thesis format.
_HEADING_PREFIX: Final[str] = "Heading "
_MAX_HEADING_LEVEL: Final[int] = 6


@dataclass(frozen=True, slots=True)
class DocxPreview:
    """A document's text, and an honest account of what is missing from it."""

    text: str
    #: True when `text` stops before the document does. The UI says so; a
    #: truncated preview presented as complete is worse than no preview.
    truncated: bool
    #: Counted rather than rendered. A reader who sees "3 张表格" knows to open
    #: the file; one who sees the surrounding prose with the tables silently
    #: absent has no way to know anything is missing.
    table_count: int


def _paragraph_markdown(paragraph: Paragraph) -> str:
    """One paragraph as a line of Markdown, or "" for an empty one."""

    text = paragraph.text.strip()
    if not text:
        return ""
    # `style.name` is `str | None` in the stubs but reaches here as `Unknown`
    # through the element this Paragraph was constructed from. Narrowed once,
    # so the heading arithmetic below is checked normally.
    named = cast("object", paragraph.style.name if paragraph.style is not None else "")
    style = named if isinstance(named, str) else ""
    if style.startswith(_HEADING_PREFIX):
        suffix = style[len(_HEADING_PREFIX) :].strip()
        if suffix.isdigit():
            level = min(int(suffix), _MAX_HEADING_LEVEL)
            return f"{'#' * level} {text}"
    if style == "Title":
        return f"# {text}"
    return text


def _table_markdown(table: Table) -> str:
    """One table as a Markdown table, header row taken from the first row.

    Cell text is flattened to a single line and pipes are escaped: a newline or
    an unescaped pipe inside a cell breaks the row it sits in, which turns one
    malformed cell into a mangled table.
    """

    rows = [
        [cell.text.replace("\n", " ").replace("|", "\\|").strip() for cell in row.cells]
        for row in table.rows
    ]
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    padded = [row + [""] * (width - len(row)) for row in rows]
    header, *body = padded
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * width) + " |",
    ]
    lines.extend("| " + " | ".join(row) + " |" for row in body)
    return "\n".join(lines)


def _body_children(document: Any) -> Iterator[Any]:
    """The body's own XML children, in document order.

    The one untyped-library boundary in this module, isolated to a function so
    the extraction below stays fully checked. python-docx ships ``py.typed``
    but annotates only its high-level surface; the lxml elements underneath are
    ``Any``, and walking them is the only way to see paragraphs and tables
    interleaved as the document has them.
    """

    return cast("Iterator[Any]", document.element.body.iterchildren())


def extract_docx_preview(content: bytes) -> DocxPreview:
    """The document's text in reading order, as Markdown.

    Paragraphs and tables are walked through the body's own XML children rather
    than through ``document.paragraphs`` and ``document.tables`` separately.
    Those two properties each return their own kind in order, but reading them
    one after the other puts every table after every paragraph -- so a report
    whose table sits in section 2 shows it after the conclusion, which is not
    what the document says.

    The preflight is here rather than at the route, so that the ceiling holds
    for every caller of this function instead of for the one path that
    remembered to ask.
    """

    preflight_docx(content)
    document = Document(io.BytesIO(content))
    pieces: list[str] = []
    tables = 0
    total = 0
    truncated = False

    for child in _body_children(document):
        tag = str(cast("object", child.tag))
        if tag.endswith("}p"):
            piece = _paragraph_markdown(Paragraph(child, document))
        elif tag.endswith("}tbl"):
            tables += 1
            piece = _table_markdown(Table(child, document))
        else:
            # sectPr and friends. Structure rather than content.
            continue
        if not piece:
            continue
        if total + len(piece) > MAX_PREVIEW_CHARS:
            truncated = True
            break
        pieces.append(piece)
        total += len(piece)

    return DocxPreview(
        text="\n\n".join(pieces),
        truncated=truncated,
        table_count=tables,
    )


__all__ = [
    "MAX_PREVIEW_CHARS",
    "MAX_PREVIEW_COMPRESSION_RATIO",
    "MAX_PREVIEW_EXPANDED_BYTES",
    "MAX_PREVIEW_ZIP_ENTRIES",
    "DocxPreview",
    "DocxTooLargeError",
    "extract_docx_preview",
    "preflight_docx",
]
