"""Deterministic in-memory DOCX rendering for the Word MCP process.

The whole module is one untyped-library boundary, which is why the suppression
below is file-level rather than the per-line form used elsewhere in this
repository. python-docx ships ``py.typed`` but annotates only its high-level
surface; everything this renderer needs -- the ``w:eastAsia`` font attribute
that keeps CJK text from rendering as boxes, table borders, repeating header
rows -- is reachable only through the raw lxml elements underneath, and those
are typed as ``Any``. Three rules therefore fired on 143 of this file's lines,
and 143 inline pragmas would bury the errors worth reading.

``reportPrivateUsage`` is off for the same reason rather than a weaker one:
``paragraph._p``, ``run._r``, ``table._tbl`` and ``cell._tc`` *are* the
documented way into the XML python-docx does not wrap. They are underscored by
convention, not withheld.

Everything that reports a *structural* mistake -- a wrong argument, a missing
attribute, an unreachable branch -- still fails the gate here, and did: the
thirteen ``paragraph_format`` accesses below are narrowed with a cast at the
one place the concrete type is known, not silenced.
"""

# pyright: reportUnknownMemberType=false
# pyright: reportUnknownVariableType=false
# pyright: reportUnknownArgumentType=false
# pyright: reportPrivateUsage=false

from __future__ import annotations

import io
import zipfile
from datetime import UTC, datetime
from typing import Any, Final, cast

from docx import Document
from docx.document import Document as WordDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.styles.style import ParagraphStyle
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from agent_workbench.apps.word_mcp.contract import DocumentRequest, SimpleTable

WORD_DOCUMENT_MEDIA_TYPE: Final[str] = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
_FIXED_PACKAGE_TIME: Final[tuple[int, int, int, int, int, int]] = (
    2000,
    1,
    1,
    0,
    0,
    0,
)
_CONTENT_WIDTH_DXA: Final[int] = 9_360
_TABLE_INDENT_DXA: Final[int] = 120
_BLUE: Final[RGBColor] = RGBColor(0x2E, 0x74, 0xB5)
_DARK_BLUE: Final[RGBColor] = RGBColor(0x1F, 0x4D, 0x78)
_BLACK: Final[RGBColor] = RGBColor(0x00, 0x00, 0x00)
_MUTED: Final[RGBColor] = RGBColor(0x62, 0x68, 0x70)
_TABLE_BORDER: Final[str] = "D9DEE5"
_TABLE_HEADER_FILL: Final[str] = "F2F4F7"


def render_document(request: DocumentRequest) -> bytes:
    """Return one complete OOXML package without touching the filesystem."""

    document = Document()
    _set_core_properties(document, request)
    bullet_num_id = _configure_document(document)
    _add_memo_masthead(document, request)

    for section in request.sections:
        document.add_paragraph(section.heading, style="Heading 1")
        for text in section.paragraphs:
            document.add_paragraph(text, style="Normal")
        for text in section.bullets:
            _add_bullet(document, text, bullet_num_id)
        if section.table is not None:
            _add_table(document, section.table)

    stream = io.BytesIO()
    document.save(stream)
    return _canonicalize_package(stream.getvalue())


def _set_core_properties(document: WordDocument, request: DocumentRequest) -> None:
    fixed = datetime(*_FIXED_PACKAGE_TIME, tzinfo=UTC)
    properties = document.core_properties
    properties.title = request.title
    properties.subject = request.subtitle or "Business brief"
    properties.author = "Agent Workbench"
    properties.last_modified_by = "Agent Workbench"
    properties.created = fixed
    properties.modified = fixed
    properties.revision = 1
    properties.keywords = "agent-workbench, word, business-brief"


def _configure_document(document: WordDocument) -> int:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = document.styles
    normal = _paragraph_style(styles, "Normal")
    _set_style_font(normal, name="Calibri", size=11, color=_BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    _configure_heading(_paragraph_style(styles, "Heading 1"), 16, _BLUE, 16, 8)
    _configure_heading(_paragraph_style(styles, "Heading 2"), 13, _BLUE, 12, 6)
    _configure_heading(_paragraph_style(styles, "Heading 3"), 12, _DARK_BLUE, 8, 4)

    title = _new_paragraph_style(styles, "AW Title")
    _set_style_font(title, name="Calibri", size=23, color=_BLACK, bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True

    subtitle = _new_paragraph_style(styles, "AW Subtitle")
    _set_style_font(subtitle, name="Calibri", size=14, color=_MUTED)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle.paragraph_format.keep_with_next = True

    bullet = _new_paragraph_style(styles, "AW Bullet")
    bullet.base_style = normal
    _set_style_font(bullet, name="Calibri", size=11, color=_BLACK)
    bullet.paragraph_format.space_before = Pt(0)
    bullet.paragraph_format.space_after = Pt(8)
    bullet.paragraph_format.line_spacing = 1.167

    _configure_header(section.header.paragraphs[0])
    _configure_footer(section.footer.paragraphs[0])
    return _add_bullet_numbering(document)


def _paragraph_style(styles: Any, name: str) -> ParagraphStyle:
    """The existing paragraph style ``name``, typed as one.

    ``styles[...]`` is annotated as returning ``BaseStyle``, which has no
    ``paragraph_format``. Every style this module reaches for is a paragraph
    style, so the narrowing is a fact about the call rather than an assumption
    about the library -- and stating it here beats thirteen suppressions.
    """

    return cast("ParagraphStyle", styles[name])


def _new_paragraph_style(styles: Any, name: str) -> ParagraphStyle:
    """Add a paragraph style and keep its concrete type."""

    return cast("ParagraphStyle", styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH))


def _configure_heading(
    style: ParagraphStyle,
    size: float,
    color: RGBColor,
    before: float,
    after: float,
) -> None:
    _set_style_font(style, name="Calibri", size=size, color=color, bold=True)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(before)
    paragraph_format.space_after = Pt(after)
    paragraph_format.keep_with_next = True


def _set_style_font(
    style: ParagraphStyle,
    *,
    name: str,
    size: float,
    color: RGBColor,
    bold: bool | None = None,
) -> None:
    font = style.font
    font.name = name
    font.size = Pt(size)
    font.color.rgb = color
    if bold is not None:
        font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _configure_header(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("AGENT WORKBENCH  ·  WORD ARTIFACT")
    _set_run_font(run, name="Calibri", size=8.5, color=_MUTED, bold=True)


def _configure_footer(paragraph: Paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    label = paragraph.add_run("Page ")
    _set_run_font(label, name="Calibri", size=8.5, color=_MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    label._r.addnext(begin)
    begin.addnext(instruction)
    instruction.addnext(end)


def _add_memo_masthead(document: WordDocument, request: DocumentRequest) -> None:
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_before = Pt(12)
    kicker.paragraph_format.space_after = Pt(5)
    run = kicker.add_run("BUSINESS BRIEF")
    _set_run_font(run, name="Calibri", size=9, color=_BLUE, bold=True)

    document.add_paragraph(request.title, style="AW Title")
    if request.subtitle is not None:
        document.add_paragraph(request.subtitle, style="AW Subtitle")

    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(2)
    rule.paragraph_format.space_after = Pt(12)
    _set_bottom_border(rule, color="2E74B5", size=12)


def _set_bottom_border(paragraph: Paragraph, *, color: str, size: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def _add_bullet_numbering(document: WordDocument) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(value)
        for node in numbering.findall(qn("w:abstractNum"))
        if (value := node.get(qn("w:abstractNumId"))) is not None
    ]
    num_ids = [
        int(value)
        for node in numbering.findall(qn("w:num"))
        if (value := node.get(qn("w:numId"))) is not None
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "bullet")
    level.append(number_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.append(indent)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    instance = OxmlElement("w:num")
    instance.set(qn("w:numId"), str(num_id))
    reference = OxmlElement("w:abstractNumId")
    reference.set(qn("w:val"), str(abstract_id))
    instance.append(reference)
    numbering.append(instance)
    return num_id


def _add_bullet(document: WordDocument, text: str, num_id: int) -> None:
    paragraph = document.add_paragraph(text, style="AW Bullet")
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.append(level)
    num_pr.append(number)


def _add_table(document: WordDocument, value: SimpleTable) -> None:
    table = document.add_table(rows=1, cols=len(value.headers))
    table.style = "Table Grid"
    widths = _column_widths(value)
    _set_table_geometry(table, widths)
    _write_row(table.rows[0].cells, value.headers, widths, header=True)
    _repeat_as_header(table.rows[0]._tr)
    for row in value.rows:
        _write_row(table.add_row().cells, row, widths, header=False)
    trailing = document.add_paragraph()
    trailing.paragraph_format.space_before = Pt(2)
    trailing.paragraph_format.space_after = Pt(0)


def _column_widths(value: SimpleTable) -> tuple[int, ...]:
    weights: list[int] = []
    for index, header in enumerate(value.headers):
        column = [header, *(row[index] for row in value.rows)]
        # Headers deserve a little extra width; very long cells wrap instead of
        # taking the whole page from their neighbours.
        weights.append(max(8, min(48, max(len(text) for text in column))))
    total = sum(weights)
    widths = [(_CONTENT_WIDTH_DXA * weight) // total for weight in weights]
    widths[-1] += _CONTENT_WIDTH_DXA - sum(widths)
    return tuple(widths)


def _set_table_geometry(table: Table, widths: tuple[int, ...]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_width = tbl_pr.first_child_found_in("w:tblW")
    tbl_width.set(qn("w:type"), "dxa")
    tbl_width.set(qn("w:w"), str(_CONTENT_WIDTH_DXA))
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), str(_TABLE_INDENT_DXA))
    tbl_pr.append(indent)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    _set_table_margins(tbl_pr)
    _set_table_borders(tbl_pr)

    for grid_column, width in zip(table._tbl.tblGrid.gridCol_lst, widths, strict=True):
        grid_column.set(qn("w:w"), str(width))


def _set_table_margins(tbl_pr: object) -> None:
    margins = OxmlElement("w:tblCellMar")
    for side, width in (("top", 80), ("left", 120), ("bottom", 80), ("right", 120)):
        element = OxmlElement(f"w:{side}")
        element.set(qn("w:w"), str(width))
        element.set(qn("w:type"), "dxa")
        margins.append(element)
    tbl_pr.append(margins)  # type: ignore[attr-defined]


def _set_table_borders(tbl_pr: object) -> None:
    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge = OxmlElement(f"w:{name}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "0")
        edge.set(qn("w:color"), _TABLE_BORDER)
        borders.append(edge)
    tbl_pr.append(borders)  # type: ignore[attr-defined]


def _write_row(
    cells: tuple[_Cell, ...],
    values: tuple[str, ...],
    widths: tuple[int, ...],
    *,
    header: bool,
) -> None:
    for cell, text, width in zip(cells, values, widths, strict=True):
        _set_cell_width(cell, width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if header:
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), _TABLE_HEADER_FILL)
            cell._tc.get_or_add_tcPr().append(shading)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(text)
        _set_run_font(
            run,
            name="Calibri",
            size=10.5,
            color=_BLACK,
            bold=header,
        )


def _set_cell_width(cell: _Cell, width: int) -> None:
    tc_width = cell._tc.get_or_add_tcPr().get_or_add_tcW()
    tc_width.set(qn("w:type"), "dxa")
    tc_width.set(qn("w:w"), str(width))


def _repeat_as_header(row: object) -> None:
    properties = row.get_or_add_trPr()  # type: ignore[attr-defined]
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def _set_run_font(
    run: Run,
    *,
    name: str,
    size: float,
    color: RGBColor,
    bold: bool | None = None,
) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _canonicalize_package(content: bytes) -> bytes:
    """Remove ZIP writer clocks and ordering as retry-sensitive entropy."""

    source_stream = io.BytesIO(content)
    target_stream = io.BytesIO()
    with (
        zipfile.ZipFile(source_stream, mode="r") as source,
        zipfile.ZipFile(
            target_stream,
            mode="w",
            compression=zipfile.ZIP_DEFLATED,
            compresslevel=9,
        ) as target,
    ):
        for item in sorted(source.infolist(), key=lambda entry: entry.filename):
            info = zipfile.ZipInfo(item.filename, date_time=_FIXED_PACKAGE_TIME)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.create_system = 0
            info.external_attr = 0
            target.writestr(info, source.read(item.filename), compresslevel=9)
    return target_stream.getvalue()


__all__ = ["WORD_DOCUMENT_MEDIA_TYPE", "render_document"]
