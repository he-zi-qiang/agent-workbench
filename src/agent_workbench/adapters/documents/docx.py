"""Reading a .docx back as text, so the UI can show one instead of describing it.

A rendered document was the one Task output the console could not display. Text
and JSON previewed inline; a .docx -- the artifact most Tasks are actually asked
for -- rendered as "这个类型只能下载查看", which tells the reader everything
about the file except what it says.

Extraction lives on the server rather than in the browser, and the reason is not
convenience. A .docx is a zip of XML, so previewing one client-side means
shipping a zip reader and an XML parser to every page load in order to re-derive
text this process can already produce -- ``python-docx`` is a core dependency
because the Word MCP server renders with it. The same library reads.

**This is a preview, not a conversion.** It recovers the document's text in
reading order and drops everything that made it a document: styles, images,
headers, page geometry. That is the correct trade for a panel beside the run --
somebody who needs the document downloads it, and the download is unchanged
bytes. What this must never do is look like more than it is, which is why the
response says how much it dropped rather than presenting a silent excerpt as the
whole file.

"How much" is a count per kind rather than a flag, and the difference matters
because omissions are not equally guessable. Truncated prose announces itself:
the text stops mid-sentence. A dropped chart does not -- the paragraphs around
it read as a complete argument, so a reader is never prompted to wonder what the
figure showed. The counts below exist so that the panel can say the document
holds four pictures and two footnotes the panel does not, which is the only form
in which that fact reaches anybody: it is not in the event stream, not in the
``ToolResult``, and not in the artifact's metadata.
"""

from __future__ import annotations

import io
import zipfile
from collections.abc import Iterator
from dataclasses import dataclass
from typing import Any, Final, cast

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

#: Bounded for the same reason every other preview here is: this is one panel
#: beside a run, and a reader who wants a 200-page report wants the file.
MAX_PREVIEW_CHARS: Final[int] = 40_000

#: Entries in the package. A .docx has on the order of ten to thirty; a couple
#: of hundred means an image-heavy one. Far past that is not a document, and
#: refusing on the count costs nothing because the count is in the central
#: directory -- no member has to be decompressed to read it.
MAX_PREVIEW_ZIP_ENTRIES: Final[int] = 512

#: What the package may weigh once opened. The caller's ceiling is on the
#: *compressed* bytes, which is not the cost of reading one: a .docx is a zip,
#: python-docx expands the whole package, and XML compresses roughly ten to
#: twenty times. So a file inside a 20 MiB limit can still ask this process for
#: a multi-hundred-megabyte allocation, which is the gap this closes.
MAX_PREVIEW_EXPANDED_BYTES: Final[int] = 100 * 1024 * 1024

#: Expanded-to-stored ratio, refused above this. Catches the small file that is
#: nothing but expansion -- a few hundred kilobytes that unpack into tens of
#: megabytes pass the absolute ceiling above while still being a bomb. Set well
#: clear of ordinary documents: prose XML lands around ten to twenty, and the
#: thesis documents this project renders measure in that range.
MAX_PREVIEW_COMPRESSION_RATIO: Final[int] = 200


class DocxTooLargeError(ValueError):
    """The package would cost more to open than a preview may spend.

    Separate from a parse failure because the two mean different things to a
    caller: a document that will not parse is a fact about the file, while this
    is a refusal to try. The route maps them to different statuses.
    """


def preflight_docx(content: bytes) -> None:
    """Refuse a package whose expansion this process should not attempt.

    Runs before ``Document`` sees the bytes, because by then it is too late --
    python-docx expands the package to parse it, so a check that ran after it
    would be reporting an allocation that already happened.

    **Why the declared sizes are enough here.** ``ZipInfo.file_size`` is read
    out of the archive's own central directory, so on its face it is a claim
    the file makes about itself, and the obvious worry is a bomb that simply
    understates it. It cannot: ``ZipExtFile`` stops a member at exactly
    ``file_size`` bytes and then fails the CRC, so an understated declaration
    yields ``BadZipFile`` rather than more bytes. python-docx reads through the
    same ``zipfile``, so whatever bounds this bounds it. An earlier draft of
    this function decompressed every member to measure it instead; that pass
    could not observe anything the declaration did not already permit, and
    charged every legitimate preview a full extra expansion for it.
    """

    try:
        archive = zipfile.ZipFile(io.BytesIO(content))
    except zipfile.BadZipFile as error:
        # Not a size refusal: this one is a parse failure, and is raised as the
        # kind of error the caller already handles for unreadable content.
        raise ValueError("the file is not a readable .docx package") from error

    with archive:
        entries = archive.infolist()
        if len(entries) > MAX_PREVIEW_ZIP_ENTRIES:
            raise DocxTooLargeError(
                f"the package holds {len(entries)} entries, "
                f"more than the {MAX_PREVIEW_ZIP_ENTRIES} a preview will open"
            )
        expanded = sum(entry.file_size for entry in entries)

    if expanded > MAX_PREVIEW_EXPANDED_BYTES:
        raise DocxTooLargeError(
            f"the package expands to {expanded} bytes, "
            f"more than the {MAX_PREVIEW_EXPANDED_BYTES} a preview will hold"
        )
    # Against the whole stored object rather than per member, so that splitting
    # one bomb across several entries does not evade it.
    if content and expanded > len(content) * MAX_PREVIEW_COMPRESSION_RATIO:
        raise DocxTooLargeError(
            f"the package expands {expanded // len(content)}x, "
            f"past the {MAX_PREVIEW_COMPRESSION_RATIO}x a preview will open"
        )


#: How a heading becomes Markdown. python-docx reports built-in heading styles
#: as "Heading 1".."Heading 9"; anything else is a paragraph, including the
#: custom styles the project's own renderer applies for Chinese thesis format.
#: That flattening is not silent -- ``flattened_paragraph_count`` reports how
#: many paragraphs it happened to.
_HEADING_PREFIX: Final[str] = "Heading "
_MAX_HEADING_LEVEL: Final[int] = 6

#: Marks a style the document defined rather than one Word ships. OOXML records
#: this itself, on ``w:style``, which is why the judgement below reads the flag
#: instead of matching names.
#:
#: An earlier version of this count matched a four-name allowlist of built-ins
#: -- "Normal", "Body Text", "No Spacing", "Plain Text" -- and counted anything
#: else as flattened. It was wrong on real documents in a way this project's
#: own fixtures could not show, because every fixture here is either
#: ``render_document``'s output or a python-docx default, and both wear
#: "Normal". A document written in Word wears "List Paragraph" on every list
#: item, "Body Text 2" on indented prose, "Quote" and "Caption" where those
#: were used -- all built-in, none of them custom. A twenty-item list reported
#: twenty flattened paragraphs *and* twenty numbered ones: two numbers, in one
#: panel, describing the same twenty paragraphs.
_WML_NAMESPACE: Final[str] = (
    "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
)
_CUSTOM_STYLE_ATTRIBUTE: Final[str] = f"{{{_WML_NAMESPACE}}}customStyle"
_STYLE_ID_ATTRIBUTE: Final[str] = f"{{{_WML_NAMESPACE}}}styleId"

#: Where the pictures are. ``document.inline_shapes`` is the obvious way to
#: write this count and is the wrong one: its XPath matches ``wp:inline`` under
#: a run and nothing else, so a document laid out with floating images -- which
#: is how most people place a figure they want text to wrap around -- reports
#: no pictures at all. A zero that means "none" and a zero that means "none of
#: the kind I looked for" are indistinguishable on screen, which is the whole
#: failure this count exists to prevent.
#:
#: VML pictures (``w:pict``, Word 2003 and some converters) are deliberately
#: not counted: that element also carries text boxes and drawn lines, so
#: including it would report shapes that are not images, and a count inflated
#: by furniture is worse than a count that is honest about its two formats.
_IMAGE_PATHS: Final[tuple[str, ...]] = (".//wp:inline", ".//wp:anchor")

#: One reference per header or footer a section defines. Equivalent to asking
#: every section's six header/footer objects whether ``is_linked_to_previous``
#: is false, because that property *is* the absence of this element -- and
#: reachable without the guard the property route needs, since a package
#: assembled by a converter can arrive with no ``sectPr`` and therefore no
#: sections to iterate.
_HEADER_PATH: Final[str] = ".//w:headerReference"
_FOOTER_PATH: Final[str] = ".//w:footerReference"

#: Footnote marks in the body, which is where the reader loses something. The
#: note text lives in a separate part; counting the definitions there would
#: also count the two ``w:type`` separators Word writes into every footnotes
#: part whether or not the author used one.
_FOOTNOTE_PATH: Final[str] = ".//w:footnoteReference"

#: Paragraphs Word numbers for itself. ``w:numId`` of 0 is excluded because
#: OOXML spends it on the opposite meaning -- it is how a paragraph cancels the
#: numbering its style would otherwise give it, so counting it would report a
#: list item where the document went out of its way to say there is none.
#:
#: Numbering carried by a style rather than by the paragraph is invisible here:
#: a "List Number" paragraph has no ``w:numPr`` of its own, and following it
#: would mean walking the ``basedOn`` chain in ``styles.xml``. Under-reporting
#: is the direction this errs in, and it is the safe one only because the
#: alternative was over-reporting every paragraph of a document that defines
#: such a style.
_NUMBERED_PARAGRAPH_PATH: Final[str] = './/w:p[w:pPr/w:numPr/w:numId[@w:val!="0"]]'


@dataclass(frozen=True, slots=True)
class DocxPreview:
    """A document's text, and an honest account of what is missing from it."""

    text: str
    #: True when `text` stops before the document does. The UI says so; a
    #: truncated preview presented as complete is worse than no preview.
    truncated: bool
    #: Counted rather than rendered. A reader who sees "3 张表格" knows to open
    #: the file; one who sees the surrounding prose with the tables silently
    #: absent has no way to know anything is missing.
    #:
    #: The one count of what the *walk reached* rather than of the document: a
    #: truncated preview reports the tables above the cut. Every count below it
    #: is of the whole document. The seam is left where it is because this
    #: number is already on screen -- the console prints 共 N 张表格 from it --
    #: and moving it is a change to what the UI claims, which nothing else here
    #: is asking for.
    table_count: int
    #: Pictures, counted rather than shown, and the sharpest case for counting
    #: anything at all: prose that reads around a figure still reads as a
    #: finished argument, so an absent image is the one omission a reader
    #: cannot infer from what is on screen.
    image_count: int
    #: Header definitions this reader never opens. They live in their own parts
    #: of the package, and folding them in would interleave a running title and
    #: a page number with the prose -- text the document shows on every page
    #: and means once. Counted so that "the letterhead is gone" is something
    #: the reader is told rather than something they find in the download.
    header_count: int
    #: The same at the other end of the page, and a separate number because a
    #: document may define one without the other.
    footer_count: int
    #: Paragraphs Word numbers for itself. The digits are not in the text --
    #: Word generates them at layout time -- so an ordered list previews as
    #: unordered lines, and a procedure whose steps had an order arrives
    #: without one. Rendering the numbers instead would mean reimplementing
    #: Word's list counter, restart rules and level inheritance included, which
    #: is a layout engine rather than a preview.
    numbered_paragraph_count: int
    #: Footnote marks. What the reader loses is not the superscript but the
    #: sentence it pointed at, which is usually the qualification on the claim
    #: the paragraph makes.
    footnote_count: int
    #: Paragraphs whose words came through and whose structure did not. The
    #: Markdown above is derived from built-in style names alone -- "Heading
    #: 1".."Heading 9" and "Title" -- so a paragraph wearing any other named
    #: style is emitted as a bare line, including every style this project's
    #: own renderer applies for Chinese thesis format. That limit was recorded
    #: in a comment next to the code that has it, where nobody reading a
    #: preview could see it; this is the same fact, in the result.
    #:
    #: Body-text styles are excluded, or an ordinary document would report
    #: every paragraph it has, and a number that tracks length says nothing
    #: about what was lost. Paragraphs inside table cells are excluded for the
    #: other reason: their styles go to a Markdown table by design, not by this
    #: reader running out of vocabulary.
    flattened_paragraph_count: int


def _style_name(paragraph: Paragraph) -> str:
    """The paragraph's style name, or "" when it has none.

    `style.name` is `str | None` in the stubs but reaches here as `Unknown`
    through the element this Paragraph was constructed from. Narrowed once, in
    one place, so that both the heading arithmetic and the count of what that
    arithmetic could not express are checked normally.
    """

    named = cast("object", paragraph.style.name if paragraph.style is not None else "")
    return named if isinstance(named, str) else ""


def _heading_level(style: str) -> int | None:
    """The heading level this style asks for, or None if this reader has none.

    Split out of ``_paragraph_markdown`` because "no level" is the definition
    of a flattened paragraph, and a second implementation of that judgement --
    a list of style names, say -- would drift from the one that renders. The
    count then reports a loss the renderer did not have, or misses one it did.
    """

    if style.startswith(_HEADING_PREFIX):
        suffix = style[len(_HEADING_PREFIX) :].strip()
        if suffix.isdigit():
            return min(int(suffix), _MAX_HEADING_LEVEL)
    if style == "Title":
        return 1
    return None


def _paragraph_markdown(paragraph: Paragraph) -> str:
    """One paragraph as a line of Markdown, or "" for an empty one."""

    text = paragraph.text.strip()
    if not text:
        return ""
    level = _heading_level(_style_name(paragraph))
    if level is None:
        return text
    return f"{'#' * level} {text}"


def _table_markdown(table: Table) -> str:
    """One table as a Markdown table, header row taken from the first row.

    Cell text is flattened to a single line and pipes are escaped: a newline or
    an unescaped pipe inside a cell breaks the row it sits in, which turns one
    malformed cell into a mangled table.
    """

    rows = [
        [cell.text.replace("\n", " ").replace("|", "\\|").strip() for cell in row.cells]
        for row in table.rows
    ]
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    padded = [row + [""] * (width - len(row)) for row in rows]
    header, *body = padded
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * width) + " |",
    ]
    lines.extend("| " + " | ".join(row) + " |" for row in body)
    return "\n".join(lines)


def _body_children(document: Any) -> Iterator[Any]:
    """The body's own XML children, in document order.

    The one untyped-library boundary in this module, isolated to a function so
    the extraction below stays fully checked. python-docx ships ``py.typed``
    but annotates only its high-level surface; the lxml elements underneath are
    ``Any``, and walking them is the only way to see paragraphs and tables
    interleaved as the document has them.
    """

    return cast("Iterator[Any]", document.element.body.iterchildren())


def _count_in_body(document: Any, *paths: str) -> int:
    """How many elements the body holds on these paths, over the whole document.

    The second untyped boundary here, isolated for the same reason as the one
    above: ``xpath`` on a python-docx element returns ``Any``, and the counts
    that call this are ordinary integers afterwards.

    Whole-document by construction rather than by care. Counting inside the
    walk that builds the text would tie every number to where that walk
    stopped, so a truncated preview would report the pictures above the cut and
    nothing about the ones below -- and the reader would take "truncated, one
    picture" to mean the remainder is prose.
    """

    body: Any = document.element.body
    return sum(len(cast("list[Any]", body.xpath(path))) for path in paths)


def _custom_style_ids(document: Any) -> frozenset[str]:
    """The ids of styles this document defined for itself.

    Read from the styles part once, because the alternative -- deciding per
    paragraph whether a style name looks built-in -- is the guess that made the
    earlier version of this count wrong (see ``_CUSTOM_STYLE_ATTRIBUTE``).
    """

    styles = cast("Iterator[Any]", document.styles.element.iterchildren())
    found: set[str] = set()
    for style in styles:
        if not str(cast("object", style.tag)).endswith("}style"):
            continue
        if str(cast("object", style.get(_CUSTOM_STYLE_ATTRIBUTE))) not in {"1", "true"}:
            continue
        style_id = cast("object", style.get(_STYLE_ID_ATTRIBUTE))
        if isinstance(style_id, str):
            found.add(style_id)
    return frozenset(found)


def _count_flattened_paragraphs(document: Any) -> int:
    """Paragraphs wearing a style this document defined and this preview drops.

    Walks the body's own children rather than ``document.paragraphs`` so that
    it sees exactly the paragraphs ``_paragraph_markdown`` decides on -- the
    body's direct children, table cells excluded -- and judges them with the
    same ``_heading_level``. Two ways of asking the same question would be two
    things to keep in step.

    **Custom styles only**, and that is a narrowing rather than a shortcut. A
    built-in style the preview also drops -- "Quote", "Caption" -- is a real
    loss and is not reported here, which is the cost. What it buys is a number
    that means what it says: this document defined a style, and the preview
    could not express it. The alternative counted every list item in every
    document Word wrote, alongside the numbering count that had already
    reported those same paragraphs.

    An empty paragraph is not counted: nothing of it reached the text, so
    nothing about it was flattened. Neither is one with no style at all, which
    is a paragraph that never asked to be anything.
    """

    # Style names are resolved once per style id rather than once per
    # paragraph. `paragraph.style` searches the styles part on every access,
    # and that search is the whole cost of this pass: 137ms for a
    # 663-paragraph document, against 10ms for all of the same document's
    # text. Ingestion calls this with `max_chars=None`, so it is the whole file
    # every time, which is where a per-paragraph lookup would be paid. Caching
    # on the id off the paragraph's own XML is exact rather than approximate --
    # one document resolves one id to one style.
    custom = _custom_style_ids(document)
    if not custom:
        return 0
    names: dict[str | None, str] = {}
    flattened = 0
    for child in _body_children(document):
        if not str(cast("object", child.tag)).endswith("}p"):
            continue
        paragraph = Paragraph(child, document)
        if not paragraph.text.strip():
            continue
        declared = cast("object", child.style)
        style_id = declared if isinstance(declared, str) else None
        if style_id is None or style_id not in custom:
            continue
        if style_id not in names:
            names[style_id] = _style_name(paragraph)
        if _heading_level(names[style_id]) is None:
            flattened += 1
    return flattened


def extract_docx_preview(
    content: bytes, *, max_chars: int | None = MAX_PREVIEW_CHARS
) -> DocxPreview:
    """The document's text in reading order, as Markdown.

    ``max_chars`` is the caller's ceiling rather than this module's, because
    the two callers want different answers and only one of them wants a
    preview. The panel beside a run stops at ``MAX_PREVIEW_CHARS`` -- a reader
    who wants a 200-page report wants the file. Ingestion passes ``None``: a
    document is about to be chunked and indexed, and stopping at forty thousand
    characters would index a fraction of it while every layer reported success.
    That is the failure this parameter exists to make impossible to reach by
    accident, which is also why the default is the preview number: the caller
    that needs the whole document has to say so.

    Paragraphs and tables are walked through the body's own XML children rather
    than through ``document.paragraphs`` and ``document.tables`` separately.
    Those two properties each return their own kind in order, but reading them
    one after the other puts every table after every paragraph -- so a report
    whose table sits in section 2 shows it after the conclusion, which is not
    what the document says.

    The preflight is here rather than at the route, so that the ceiling holds
    for every caller of this function instead of for the one path that
    remembered to ask.

    The counts are taken from the whole body after that walk has stopped, not
    accumulated during it, so a preview cut at ``max_chars`` still reports the
    pictures below the cut. ``table_count`` is the exception and keeps counting
    what the walk reached; see the field.
    """

    preflight_docx(content)
    document = Document(io.BytesIO(content))
    pieces: list[str] = []
    tables = 0
    total = 0
    truncated = False

    for child in _body_children(document):
        tag = str(cast("object", child.tag))
        if tag.endswith("}p"):
            piece = _paragraph_markdown(Paragraph(child, document))
        elif tag.endswith("}tbl"):
            tables += 1
            piece = _table_markdown(Table(child, document))
        else:
            # sectPr and friends. Structure rather than content.
            continue
        if not piece:
            continue
        if max_chars is not None and total + len(piece) > max_chars:
            truncated = True
            break
        pieces.append(piece)
        total += len(piece)

    return DocxPreview(
        text="\n\n".join(pieces),
        truncated=truncated,
        table_count=tables,
        image_count=_count_in_body(document, *_IMAGE_PATHS),
        header_count=_count_in_body(document, _HEADER_PATH),
        footer_count=_count_in_body(document, _FOOTER_PATH),
        numbered_paragraph_count=_count_in_body(document, _NUMBERED_PARAGRAPH_PATH),
        footnote_count=_count_in_body(document, _FOOTNOTE_PATH),
        flattened_paragraph_count=_count_flattened_paragraphs(document),
    )


__all__ = [
    "MAX_PREVIEW_CHARS",
    "MAX_PREVIEW_COMPRESSION_RATIO",
    "MAX_PREVIEW_EXPANDED_BYTES",
    "MAX_PREVIEW_ZIP_ENTRIES",
    "DocxPreview",
    "DocxTooLargeError",
    "extract_docx_preview",
    "preflight_docx",
]
